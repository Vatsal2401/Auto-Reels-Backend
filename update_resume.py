import docx
from docx.shared import Pt

source_file = '/home/vatsal2401/Downloads/Vatsal_Patel_Resume.docx'
target_file = '/home/vatsal2401/Downloads/Vatsal_Patel_Resume_Updated.docx'

doc = docx.Document(source_file)

# We want to replace "Jul 2022" with "Oct 2023" for Saleshandy
# and insert Rapidops experience before "PROJECTS"

projects_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Jul 2022" in p.text and "Saleshandy" in p.text:
        # Update paragraph text by modifying runs to preserve some formatting if possible
        for run in p.runs:
            if "Jul 2022" in run.text:
                run.text = run.text.replace("Jul 2022", "Oct 2023")
                break
        else:
            # Fallback if text is split across runs
            p.text = p.text.replace("Jul 2022", "Oct 2023")

    if "PROJECTS" in p.text:
        projects_idx = i
        break

if projects_idx != -1:
    # Insert new experience before projects
    p_projects = doc.paragraphs[projects_idx]
    
    # Try to find a bullet style by looking at previous paragraphs
    bullet_style = None
    for j in range(projects_idx-1, -1, -1):
        if doc.paragraphs[j].style and 'List' in doc.paragraphs[j].style.name:
            bullet_style = doc.paragraphs[j].style
            break
            
    # Find heading style
    heading_style = None
    for j in range(projects_idx-1, -1, -1):
        if "Saleshandy" in doc.paragraphs[j].text:
            heading_style = doc.paragraphs[j].style
            break
            
    # Add Rapidops Heading
    new_p1 = p_projects.insert_paragraph_before("", style=heading_style)
    # We'll construct styling similar to Saleshandy manually if we need to, but let's just make it bold
    r1 = new_p1.add_run("Software Engineer Intern  |  Rapidops Inc. ")
    r1.bold = True
    r2 = new_p1.add_run(" Ahmedabad, IN ")
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    # If the original has a specific format like tab stops, we just add the date
    r3 = new_p1.add_run("Jan 2023 – Jul 2023")
    r3.bold = False

    # Add bullets
    b1 = p_projects.insert_paragraph_before("Developed and maintained microservices utilizing Apache Kafka for asynchronous event-driven architecture, improving data stream reliability.", style=bullet_style or 'List Paragraph')
    b2 = p_projects.insert_paragraph_before("Built interactive and responsive frontend components utilizing React.js to streamline user workflows and visualize pipeline data.", style=bullet_style or 'List Paragraph')
    b3 = p_projects.insert_paragraph_before("Collaborated across teams to integrate new features and optimize system performance during the internship period.", style=bullet_style or 'List Paragraph')
    
    # Insert a blank line if needed
    p_projects.insert_paragraph_before()

doc.save(target_file)
print(f"Updated resume saved to {target_file}")
